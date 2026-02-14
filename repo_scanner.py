#!/usr/bin/env python3
"""
Repository Dump Generator - Word & PDF Export
----------------------------------------------
Recursively scans a repository and creates a formatted document containing:
- Complete folder structure (tree format)
- Full content of all text-based files
- Excludes binary and media files
- Outputs to Word (.docx) or PDF format

Requirements:
    pip install python-docx reportlab

Author: Python Repository Scanner
Compatible with: Windows, macOS, Linux
Python Version: 3.6+
"""

import os
import re
from pathlib import Path
from typing import List, Set
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Preformatted
from reportlab.lib.enums import TA_LEFT

# Configuration: Text file extensions to include
TEXT_EXTENSIONS = {
    '.py', '.js', '.jsx', '.ts', '.tsx', '.html', '.css', '.scss', '.sass',
    '.json', '.xml', '.yaml', '.yml', '.txt', '.md', '.markdown', '.rst',
    '.c', '.cpp', '.h', '.hpp', '.java', '.go', '.rs', '.php', '.rb',
    '.sh', '.bash', '.sql', '.r', '.m', '.swift', '.kt', '.scala',
    '.vue', '.svelte', '.conf', '.config', '.ini', '.toml', '.env',
    '.log', '.csv', '.tsv', '.gitignore', '.dockerignore', 'Dockerfile',
    'Makefile', '.editorconfig', 'requirements.txt', 'package.json'
}

# Configuration: Binary/media extensions to exclude
BINARY_EXTENSIONS = {
    '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.svg', '.webp',
    '.mp4', '.avi', '.mov', '.mkv', '.mp3', '.wav', '.flac',
    '.zip', '.tar', '.gz', '.rar', '.7z', '.bz2',
    '.exe', '.dll', '.so', '.dylib', '.bin',
    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
    '.pyc', '.pyo', '.class', '.o', '.obj',
    '.db', '.sqlite', '.sqlite3'
}

# Configuration: Directories to skip
SKIP_DIRECTORIES = {
    '.git', '.svn', '.hg', '__pycache__', 'node_modules',
    '.venv', 'venv', 'env', '.env', 'virtualenv',
    '.idea', '.vscode', '.vs', '.pytest_cache', '.mypy_cache',
    'dist', 'build', '.eggs', '*.egg-info', 'target'
}

# Tree structure characters
TREE_BRANCH = '├── '
TREE_LAST = '└── '
TREE_VERTICAL = '│   '
TREE_SPACE = '    '


def sanitize_xml_content(text: str) -> str:
    """
    Remove NULL bytes and control characters that are not XML-compatible.
    Preserves newlines, tabs, and carriage returns.
    """
    if not text:
        return text
    
    # Remove NULL bytes
    text = text.replace('\x00', '')
    
    # Remove control characters except for tab (\t), newline (\n), and carriage return (\r)
    # XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    return text


def should_include_file(file_path: Path) -> bool:
    """Determine if a file should be included in the dump."""
    ext = file_path.suffix.lower()
    name = file_path.name.lower()
    
    if ext in BINARY_EXTENSIONS:
        return False
    
    if ext in TEXT_EXTENSIONS:
        return True
    
    if not ext and name in {'makefile', 'dockerfile', 'readme', 'license', 'changelog'}:
        return True
    
    try:
        if file_path.stat().st_size > 10 * 1024 * 1024:
            return False
    except OSError:
        return False
    
    return False


def should_skip_directory(dir_name: str) -> bool:
    """Check if a directory should be skipped."""
    return dir_name in SKIP_DIRECTORIES or dir_name.startswith('.')


def generate_tree_structure(root_path: Path, prefix: str = '', 
                           is_last: bool = True, output_lines: List[str] = None) -> List[str]:
    """Generate a tree structure of the directory."""
    if output_lines is None:
        output_lines = []
    
    try:
        items = sorted(root_path.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
        items = [item for item in items if not (item.is_dir() and should_skip_directory(item.name))]
        
        for index, item in enumerate(items):
            is_last_item = (index == len(items) - 1)
            connector = TREE_LAST if is_last_item else TREE_BRANCH
            
            if item.is_dir():
                output_lines.append(f"{prefix}{connector}{item.name}/")
            else:
                output_lines.append(f"{prefix}{connector}{item.name}")
            
            if item.is_dir():
                extension = TREE_SPACE if is_last_item else TREE_VERTICAL
                generate_tree_structure(item, prefix + extension, is_last_item, output_lines)
    
    except PermissionError:
        output_lines.append(f"{prefix}[Permission Denied]")
    except Exception as e:
        output_lines.append(f"{prefix}[Error: {str(e)}]")
    
    return output_lines


def read_file_content(file_path: Path) -> str:
    """Read file content with proper encoding handling."""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='strict') as f:
                content = f.read()
                # Sanitize the content before returning
                return sanitize_xml_content(content)
        except UnicodeDecodeError:
            continue
        except Exception as e:
            return f"[Error reading file: {str(e)}]"
    
    return "[Error: Unable to decode file with supported encodings]"


def collect_file_contents(root_path: Path, base_path: Path = None) -> List[tuple]:
    """Recursively collect all file contents from the repository."""
    if base_path is None:
        base_path = root_path
    
    file_contents = []
    
    try:
        for item in sorted(root_path.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower())):
            if item.is_dir() and should_skip_directory(item.name):
                continue
            
            if item.is_file():
                if should_include_file(item):
                    relative_path = item.relative_to(base_path)
                    content = read_file_content(item)
                    file_contents.append((str(relative_path), content))
            
            elif item.is_dir():
                file_contents.extend(collect_file_contents(item, base_path))
    
    except PermissionError:
        pass
    except Exception:
        pass
    
    return file_contents


def generate_word_document(output_file: str = 'repo_dump.docx'):
    """Generate repository dump as a Word document."""
    print("🔍 Starting repository scan for Word document...")
    
    current_dir = Path.cwd()
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(9)
    
    # Title
    title = doc.add_heading('Repository Dump', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Root Directory: {current_dir}")
    doc.add_paragraph("=" * 60)
    doc.add_paragraph()
    
    # Directory structure
    print("📁 Generating directory tree...")
    doc.add_heading('Directory Structure', 1)
    doc.add_paragraph(f"{current_dir.name}/")
    
    tree_lines = generate_tree_structure(current_dir)
    tree_text = '\n'.join(tree_lines)
    # Sanitize tree text as well
    tree_text = sanitize_xml_content(tree_text)
    p = doc.add_paragraph(tree_text)
    p.style.font.name = 'Courier New'
    p.style.font.size = Pt(8)
    
    doc.add_page_break()
    
    # File contents
    print("📄 Collecting file contents...")
    file_contents = collect_file_contents(current_dir)
    
    doc.add_heading(f'File Contents ({len(file_contents)} files)', 1)
    
    for index, (relative_path, content) in enumerate(file_contents, 1):
        print(f"   Processing: {relative_path}")
        
        # File header
        doc.add_paragraph("─" * 60)
        file_heading = doc.add_paragraph(f"File: {relative_path}")
        file_heading.runs[0].bold = True
        file_heading.runs[0].font.color.rgb = RGBColor(0, 0, 128)
        doc.add_paragraph("─" * 60)
        
        # File content - already sanitized in read_file_content()
        try:
            content_para = doc.add_paragraph(content)
            content_para.style.font.name = 'Courier New'
            content_para.style.font.size = Pt(8)
        except Exception as e:
            # If there's still an error, add an error message instead
            error_para = doc.add_paragraph(f"[Error adding content: {str(e)}]")
            error_para.style.font.name = 'Courier New'
            error_para.style.font.size = Pt(8)
        
        doc.add_paragraph()
        
        # Add page break every 3 files to avoid huge pages
        if index % 3 == 0:
            doc.add_page_break()
    
    # Footer
    doc.add_paragraph("=" * 60)
    doc.add_paragraph(f"Total files processed: {len(file_contents)}")
    
    doc.save(output_file)
    
    print(f"\n✅ Word document completed!")
    print(f"📦 Output saved to: {output_file}")
    print(f"📊 Total files processed: {len(file_contents)}")


def generate_pdf_document(output_file: str = 'repo_dump.pdf'):
    """Generate repository dump as a PDF document."""
    print("🔍 Starting repository scan for PDF document...")
    
    current_dir = Path.cwd()
    doc = SimpleDocTemplate(output_file, pagesize=letter,
                           rightMargin=0.5*inch, leftMargin=0.5*inch,
                           topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Code', fontName='Courier', fontSize=7, leading=9))
    styles.add(ParagraphStyle(name='FileHeader', fontName='Courier-Bold', fontSize=9, 
                             textColor='navy', leading=11))
    
    # Title
    title_style = styles['Heading1']
    title_style.alignment = TA_LEFT
    elements.append(Paragraph("Repository Dump", title_style))
    elements.append(Spacer(1, 12))
    
    # Metadata
    elements.append(Paragraph(f"Root Directory: {current_dir}", styles['Normal']))
    elements.append(Paragraph("=" * 100, styles['Code']))
    elements.append(Spacer(1, 12))
    
    # Directory structure
    print("📁 Generating directory tree...")
    elements.append(Paragraph("Directory Structure", styles['Heading2']))
    elements.append(Spacer(1, 6))
    
    tree_lines = generate_tree_structure(current_dir)
    tree_text = f"{current_dir.name}/\n" + '\n'.join(tree_lines)
    tree_text = sanitize_xml_content(tree_text)
    
    # Use Preformatted for tree structure to preserve spacing
    elements.append(Preformatted(tree_text, styles['Code']))
    elements.append(PageBreak())
    
    # File contents
    print("📄 Collecting file contents...")
    file_contents = collect_file_contents(current_dir)
    
    elements.append(Paragraph(f"File Contents ({len(file_contents)} files)", styles['Heading2']))
    elements.append(Spacer(1, 12))
    
    for index, (relative_path, content) in enumerate(file_contents, 1):
        print(f"   Processing: {relative_path}")
        
        # File header
        elements.append(Paragraph("─" * 100, styles['Code']))
        elements.append(Paragraph(f"File: {relative_path}", styles['FileHeader']))
        elements.append(Paragraph("─" * 100, styles['Code']))
        elements.append(Spacer(1, 6))
        
        # File content - escape special characters for PDF and sanitize
        content_escaped = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        elements.append(Preformatted(content_escaped, styles['Code']))
        elements.append(Spacer(1, 12))
        
        # Add page break every 2 files
        if index % 2 == 0:
            elements.append(PageBreak())
    
    # Footer
    elements.append(Paragraph("=" * 100, styles['Code']))
    elements.append(Paragraph(f"Total files processed: {len(file_contents)}", styles['Normal']))
    
    # Build PDF
    doc.build(elements)
    
    print(f"\n✅ PDF document completed!")
    print(f"📦 Output saved to: {output_file}")
    print(f"📊 Total files processed: {len(file_contents)}")


if __name__ == "__main__":
    try:
        print("Select output format:")
        print("1. Word Document (.docx)")
        print("2. PDF Document (.pdf)")
        print("3. Both")
        
        choice = input("\nEnter your choice (1/2/3): ").strip()
        
        if choice == '1':
            generate_word_document()
        elif choice == '2':
            generate_pdf_document()
        elif choice == '3':
            generate_word_document()
            print("\n" + "─" * 60 + "\n")
            generate_pdf_document()
        else:
            print("Invalid choice. Please run again and select 1, 2, or 3.")
    
    except KeyboardInterrupt:
        print("\n\n⚠️  Operation cancelled by user.")
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        raise